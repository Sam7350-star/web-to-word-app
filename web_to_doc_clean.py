
import streamlit as st
import requests
from bs4 import BeautifulSoup
from readability import Document
from docx import Document as DocxDocument
import tempfile
import os

st.set_page_config(page_title="Multi-Article Web to Word Exporter", layout="centered")

st.title("📰 Web Articles to Word Document")
st.write("Paste multiple article URLs below (one per line). The app will extract all the articles and generate a single downloadable Word (.docx) file.")

# Input fields
urls_input = st.text_area("Enter URLs (one per line)")
doc_title = st.text_input("Enter Word Document Title", value="Collected Articles")

if st.button("Generate Word Document"):
    urls = [u.strip() for u in urls_input.splitlines() if u.strip()]

    if not urls:
        st.warning("Please enter at least one URL.")
    else:
        try:
            doc = DocxDocument()
            doc.add_heading(doc_title, 0)

            for i, url in enumerate(urls, start=1):
                try:
                    response = requests.get(url, timeout=10)
                    response.raise_for_status()
                    readable_article = Document(response.text)
                    soup = BeautifulSoup(readable_article.summary(), 'html.parser')
                    text = soup.get_text(separator="\n").strip()

                    doc.add_heading(f"{i}. {readable_article.title()}", level=1)
                    doc.add_paragraph(text)
                    doc.add_paragraph("\n")
                except Exception as e:
                    doc.add_heading(f"{i}. [Error fetching article]", level=1)
                    doc.add_paragraph(f"URL: {url}")
                    doc.add_paragraph(f"Error: {e}")
                    doc.add_paragraph("\n")

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            with open(tmp_path, "rb") as file:
                st.download_button(
                    label="📥 Download Word Document",
                    data=file,
                    file_name=f"{doc_title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            os.remove(tmp_path)

        except Exception as e:
            st.error(f"An error occurred: {e}")
